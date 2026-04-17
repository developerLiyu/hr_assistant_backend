import os
import zipfile
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document

# ===================== 配置 =====================
OUTPUT_DIR = "resume_test_files"  # 输出文件夹
ZIP_NAME = "批量简历测试包.zip"    # 批量压缩包名
POSITION_ID = 1                    # 关联岗位ID（测试用）

# 测试简历数据（严格对齐你的Pydantic/ORM模型）
RESUME_DATA = {
    "zhangsan": {
        "name": "张三",
        "phone": "13800138000",
        "email": "zhangsan@test.com",
        "education": "本科",
        "school": "北京理工大学",
        "major": "计算机科学与技术",
        "work_years": 3,
        "company": "字节跳动",
        "position": "后端开发工程师",
        "skills": ["Python", "FastAPI", "MySQL", "Redis", "Docker"],
        "job": "后端开发",
        "suffix": "pdf"
    },
    "lisi": {
        "name": "李四",
        "phone": "13900139000",
        "email": "lisi@test.com",
        "education": "硕士",
        "school": "清华大学",
        "major": "产品设计",
        "work_years": 5,
        "company": "阿里云计算有限公司",
        "position": "高级产品经理",
        "skills": ["需求分析", "原型设计", "项目管理", "用户调研"],
        "job": "产品经理",
        "suffix": "doc"
    },
    "wangwu": {
        "name": "王五",
        "phone": "13700137000",
        "email": "wangwu@test.com",
        "education": "专科",
        "school": "深圳职业技术学院",
        "major": "前端开发",
        "work_years": 2,
        "company": "美团",
        "position": "前端开发工程师",
        "skills": ["Vue3", "React", "TypeScript", "Webpack", "小程序"],
        "job": "前端开发",
        "suffix": "docx"
    }
}

# ===================== 1. 创建输出目录 =====================
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ===================== 2. 生成PDF简历 =====================
def create_pdf(data):
    filename = f"测试简历_{data['name']}_{data['job']}.pdf"
    path = os.path.join(OUTPUT_DIR, filename)
    c = canvas.Canvas(path, pagesize=A4)
    c.setFont("Helvetica", 12)
    text = f"""候选人姓名：{data['name']}
手机号：{data['phone']}
邮箱：{data['email']}
学历：{data['education']}
毕业院校：{data['school']}
专业：{data['major']}
工作年限：{data['work_years']}
当前公司：{data['company']}
当前职位：{data['position']}

技能标签：{', '.join(data['skills'])}

工作经历：
1. 公司：{data['company']}，职位：{data['position']}，时间：2023-01 至 至今
   描述：负责核心业务开发、接口设计、数据库优化

项目经验：
1. 项目名称：简历管理系统，角色：核心开发
   描述：负责上传、解析、存储全流程开发

教育经历：
1. 学校：{data['school']}，专业：{data['major']}，学历：{data['education']}，时间：2017-09 至 2021-06

AI简历摘要：{data['work_years']}年{data['job']}经验，熟练掌握主流技术栈，业务落地能力强。
"""
    lines = text.split("\n")
    y = 800
    for line in lines:
        c.drawString(50, y, line)
        y -= 15
    c.save()
    return filename

# ===================== 3. 生成DOC简历（纯文本） =====================
def create_doc(data):
    filename = f"测试简历_{data['name']}_{data['job']}.doc"
    path = os.path.join(OUTPUT_DIR, filename)
    content = f"""候选人姓名：{data['name']}
手机号：{data['phone']}
邮箱：{data['email']}
学历：{data['education']}
毕业院校：{data['school']}
专业：{data['major']}
工作年限：{data['work_years']}
当前公司：{data['company']}
当前职位：{data['position']}

技能标签：{', '.join(data['skills'])}

工作经历：
1. 公司：{data['company']}，职位：{data['position']}，时间：2020-03 至 至今
   描述：负责产品规划、需求分析、项目落地

项目经验：
1. 项目名称：企业级云平台，角色：产品负责人
   描述：主导产品从0到1设计与迭代

教育经历：
1. 学校：{data['school']}，专业：{data['major']}，学历：{data['education']}，时间：2013-09 至 2016-06

AI简历摘要：{data['work_years']}年互联网产品经验，擅长ToB产品规划与跨部门协作。
"""
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return filename

# ===================== 4. 生成DOCX简历 =====================
def create_docx(data):
    filename = f"测试简历_{data['name']}_{data['job']}.docx"
    path = os.path.join(OUTPUT_DIR, filename)
    doc = Document()
    doc.add_paragraph(f"候选人姓名：{data['name']}")
    doc.add_paragraph(f"手机号：{data['phone']}")
    doc.add_paragraph(f"邮箱：{data['email']}")
    doc.add_paragraph(f"学历：{data['education']}")
    doc.add_paragraph(f"毕业院校：{data['school']}")
    doc.add_paragraph(f"专业：{data['major']}")
    doc.add_paragraph(f"工作年限：{data['work_years']}")
    doc.add_paragraph(f"当前公司：{data['company']}")
    doc.add_paragraph(f"当前职位：{data['position']}")
    doc.add_paragraph(f"\n技能标签：{', '.join(data['skills'])}")
    doc.add_paragraph("\n工作经历：")
    doc.add_paragraph(f"1. 公司：{data['company']}，职位：{data['position']}，时间：2022-07 至 至今")
    doc.add_paragraph("   描述：负责页面开发、性能优化、交互实现")
    doc.add_paragraph("\n项目经验：")
    doc.add_paragraph("1. 项目名称：美团商家后台，角色：前端开发")
    doc.add_paragraph("   描述：负责页面重构与用户体验优化")
    doc.add_paragraph("\n教育经历：")
    doc.add_paragraph(f"1. 学校：{data['school']}，专业：{data['major']}，学历：{data['education']}，时间：2019-09 至 2022-06")
    doc.add_paragraph(f"\nAI简历摘要：{data['work_years']}年前端开发经验，熟练掌握主流框架，擅长性能优化。")
    doc.save(path)
    return filename

# ===================== 5. 生成错误格式文件 =====================
def create_error_files():
    errors = []
    # TXT错误文件
    txt_path = os.path.join(OUTPUT_DIR, "错误格式_测试文件.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("这是不支持的文本文件")
    errors.append("错误格式_测试文件.txt")
    # JPG错误文件（空文件）
    jpg_path = os.path.join(OUTPUT_DIR, "错误格式_图片.jpg")
    with open(jpg_path, "wb") as f:
        f.write(b"")
    errors.append("错误格式_图片.jpg")
    return errors

# ===================== 6. 打包ZIP =====================
def zip_files(file_list, error_list):
    zip_path = os.path.join(OUTPUT_DIR, ZIP_NAME)
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in file_list:
            zf.write(os.path.join(OUTPUT_DIR, f), f)
        for f in error_list:
            zf.write(os.path.join(OUTPUT_DIR, f), f)
    print(f"✅ 批量ZIP包已生成：{zip_path}")

# ===================== 执行生成 =====================
if __name__ == "__main__":
    generated_files = []
    # 生成3份标准简历
    generated_files.append(create_pdf(RESUME_DATA["zhangsan"]))
    generated_files.append(create_doc(RESUME_DATA["lisi"]))
    generated_files.append(create_docx(RESUME_DATA["wangwu"]))
    # 生成错误格式文件
    error_files = create_error_files()
    # 打包ZIP
    zip_files(generated_files, error_files)
    print("\n🎉 所有测试文件生成完成！文件夹：", OUTPUT_DIR)
    print("\n📦 生成文件清单：")
    for f in generated_files + error_files + [ZIP_NAME]:
        print(f"- {f}")