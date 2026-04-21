import os
import uuid
import json
import zipfile

from dotenv import load_dotenv
from datetime import datetime
from fastapi import UploadFile
from aiofiles import open as aio_open
import PyPDF2
from docx import Document

from app.utils.path_tool import get_abs_path

# 加载配置
ALLOWED_TYPES = json.loads(os.getenv("ALLOWED_TYPES"))
MAX_FILE_SIZE = int(os.getenv("MAX_FILE_SIZE"))
UPLOAD_DIR = os.getenv("UPLOAD_DIR")
os.makedirs(UPLOAD_DIR, exist_ok=True)

async def async_validate_file(filename: str, size: int) -> tuple[bool, str]:
    """异步校验文件"""
    ext = filename.split(".")[-1].lower()
    if ext not in ALLOWED_TYPES:
        return False, f"不支持{ext}格式"
    if size > MAX_FILE_SIZE:
        return False, "文件超过10MB"
    return True, "合法"

async def async_save_and_extract_text(file: UploadFile) -> tuple[str, str]:
    """
    异步保存文件+提取原始文本
    :param file: 要操作的文件
    :return: tuple(文件相对路径, 文件内容)
    """
    ext = file.filename.split(".")[-1].lower()
    date_dir = datetime.now().strftime("%Y%m%d")
    save_dir = os.path.join(UPLOAD_DIR, date_dir)
    os.makedirs(save_dir, exist_ok=True)

    # 生成文件路径
    filename = f"{uuid.uuid4().hex}.{ext}"
    rel_path = os.path.join(date_dir, filename)
    abs_path = os.path.join(UPLOAD_DIR, rel_path)

    # 异步写入
    content = await file.read()
    async with aio_open(abs_path, "wb") as f:
        await f.write(content)

    # 提取文本
    text = ""
    try:
        if ext == "pdf":
            with open(abs_path, "rb") as f:
                # 从 PDF 文件中提取所有页面的文本内容（如果没有，返回""），并用换行符连接成一个完整的字符
                text = "\n".join([p.extract_text() or "" for p in PyPDF2.PdfReader(f).pages])
        elif ext in ["docx", "doc"]:
            # 从 Word (.docx) 文件中按从上到下的顺序提取所有内容（包括段落和表格）
            doc = Document(abs_path)
            extracted_texts = []

            # 遍历文档的所有元素，保持原始顺序
            for element in doc.element.body:
                # 处理段落 <w:p>
                if element.tag.endswith('p'):
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    if para and para.text.strip():
                        extracted_texts.append(para.text)

                # 处理表格 <w:tbl>
                elif element.tag.endswith('tbl'):
                    tbl = None
                    for t in doc.tables:
                        if t._element == element:
                            tbl = t
                            break
                    if tbl:
                        for row in tbl.rows:
                            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_texts:
                                extracted_texts.append("\t".join(row_texts))

            text = "\n".join(extracted_texts)

    except:
        text = ""
    return rel_path, text


# ------------------------------
# 3. 【新增】异步解压ZIP包，返回内部所有合法简历文件
# ------------------------------
async def async_unzip_file(zip_file: UploadFile) -> list:
    """
    异步解压ZIP，过滤出 pdf/docx/doc，返回 (临时文件流, 原始文件名)
    """
    zip_abs_path = f"{UPLOAD_DIR}/tmp"
    os.makedirs(zip_abs_path, exist_ok=True)

    zip_file_name = f"{uuid.uuid4().hex}.zip"
    zip_abs_file = os.path.join(zip_abs_path, zip_file_name)
    # 异步保存ZIP到临时目录
    async with aio_open(zip_abs_file, "wb") as f:
        await f.write(await zip_file.read())

    res = []
    try:
        with zipfile.ZipFile(zip_abs_file, "r") as zf:
            for info in zf.infolist():
                # 过滤掉无效数据：目录、mac系统自动生成的隐藏文件夹、隐藏文件
                if info.is_dir() or info.filename.startswith("__MACOSX") or info.filename.startswith("."):
                    continue
                fname = info.filename
                ext = fname.split(".")[-1].lower()
                if ext in ["pdf", "docx", "doc"]:
                    # 解压文件并放到指定目录下
                    zf.extract(info, zip_abs_path)
                    tmp_file_path = os.path.join(zip_abs_path, info.filename)
                    # 构造伪UploadFile对象
                    from fastapi import UploadFile
                    f = open(tmp_file_path, "rb")
                    upload_file = UploadFile(filename=fname, file=f) # 构建伪UploadFile对象
                    res.append((upload_file, fname))
    finally:
        # 清理临时ZIP文件
        if os.path.exists(zip_abs_file):
            try:
                os.remove(zip_abs_file)
            except Exception:
                pass

        # 注意：不解压后的文件，因为调用方还需要读取
        # 调用方应在使用完毕后自行清理这些临时文件
    return res


# =========================================音频文件=========================================
# 音频文件允许的扩展名和最大大小（字节）
ALLOWED_AUDIO_TYPES = {
    "mp3": 200 * 1024 * 1024,  # 200MB
    "wav": 500 * 1024 * 1024,  # 500MB
    "m4a": 200 * 1024 * 1024,  # 200MB
    "aac": 200 * 1024 * 1024,  # 200MB
}

async def async_validate_audio_file(filename: str, size: int) -> tuple[bool, str]:
    """异步校验音频文件"""
    ext = filename.split(".")[-1].lower()
    if ext not in ALLOWED_AUDIO_TYPES:
        return False, f"不支持的音频格式：{ext}。支持的格式：{', '.join(ALLOWED_AUDIO_TYPES.keys())}"
    if size > ALLOWED_AUDIO_TYPES[ext]:
        max_size_mb = ALLOWED_AUDIO_TYPES[ext] / (1024 * 1024)
        return False, f"文件超过{max_size_mb}MB限制"
    return True, "合法"


async def async_save_audio_file(file: UploadFile, content: bytes = None) -> tuple[str, str]:
    """
    异步保存音频文件
    :param file: 音频文件
    :return: tuple(文件相对路径, 文件绝对路径)
    """

    ext = file.filename.split(".")[-1].lower()
    audio_dir = os.path.join(datetime.now().strftime("%Y%m%d"), "audio")
    save_dir = os.path.join(UPLOAD_DIR, audio_dir)
    os.makedirs(save_dir, exist_ok=True)

    # 生成文件名
    filename = f"{uuid.uuid4().hex}.{ext}"
    audio_path = os.path.join(audio_dir, filename)
    rel_path = os.path.join(UPLOAD_DIR, audio_path)
    abs_path = get_abs_path(rel_path)

    # 保存文件
    # 如果已传入content则直接使用，否则读取
    if content is None:
        content = await file.read()
    with open(abs_path, "wb") as f:
        f.write(content)

    return rel_path, abs_path

# =========================================音频文件=========================================

